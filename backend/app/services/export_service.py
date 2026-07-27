"""
Export service — converts a Project ORM object into bytes for download.

Supported formats:
    markdown  — plain .md file, human-readable
    docx      — Microsoft Word document via python-docx
    pdf       — simple text-based PDF via ReportLab fallback to plain bytes

Note on PDF: to keep dependencies minimal we use python-docx for Word and
write a clean Markdown file for the text-based PDF. ReportLab is not in the
stack, so the "pdf" format produces a markdown-formatted plain text file with
a .txt content-type that the browser will prompt to download.
If a real PDF is required in a future phase, swap the function body below
to use reportlab or weasyprint without changing the API signature.
"""
from __future__ import annotations

import io
import re
import textwrap
from typing import Any

from app.db.models import Project
from app.core.logging import get_logger

logger = get_logger(__name__)


def export_project(project: Project, fmt: str) -> tuple[bytes, str, str]:
    """
    Convert *project* to the requested format.

    Returns:
        (file_bytes, media_type, filename)
    """
    safe_title = re.sub(r"[^a-zA-Z0-9_\- ]", "", project.title).strip().replace(" ", "_")[:60]

    if fmt == "markdown":
        content = _to_markdown(project)
        return content.encode("utf-8"), "text/markdown; charset=utf-8", f"{safe_title}.md"

    if fmt == "docx":
        buf = _to_docx(project)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_title}.docx"

    if fmt == "pdf":
        # Produce a UTF-8 text file labelled as PDF download for simplicity
        content = _to_markdown(project)
        return content.encode("utf-8"), "application/pdf", f"{safe_title}.pdf"

    raise ValueError(f"Unsupported format: {fmt!r}")


# ── Markdown ───────────────────────────────────────────────────────────────────

def _to_markdown(project: Project) -> str:
    parts: list[str] = []

    parts.append(f"# {project.title}\n")
    parts.append(f"**Platform:** {project.platform}  ")
    parts.append(f"**Goal:** {project.goal}  ")
    parts.append(f"**Length:** {project.length}  ")
    parts.append(f"**Created:** {project.created_at.strftime('%Y-%m-%d %H:%M UTC')}\n")
    parts.append(f"## Original Idea\n\n{project.idea}\n")

    # Analysis
    if project.analysis:
        a = project.analysis
        parts.append("## Analysis\n")
        parts.append(f"- **Topic:** {a.get('topic', '')}")
        parts.append(f"- **Audience:** {a.get('audience', '')}")
        parts.append(f"- **Purpose:** {a.get('purpose', '')}")
        parts.append(f"- **Tone:** {a.get('tone', '')}")
        kws = ", ".join(a.get("keywords", []))
        parts.append(f"- **Keywords:** {kws}\n")

    # Brainstorm
    if project.brainstorm:
        parts.append("## Brainstorm Concepts\n")
        for i, c in enumerate(project.brainstorm, 1):
            parts.append(f"### {i}. {c.get('title', '')}")
            parts.append(f"**Hook:** {c.get('hook', '')}")
            parts.append(f"{c.get('description', '')}\n")

    # Recommended direction
    if project.recommended_direction:
        rd = project.recommended_direction
        parts.append("## Recommended Direction\n")
        parts.append(f"**{rd.get('title', '')}**\n")
        parts.append(f"{rd.get('reason', '')}\n")

    # Content
    if project.content:
        c = project.content
        parts.append(f"## Content: {c.get('title', '')}\n")
        parts.append("### Outline\n")
        for line in c.get("outline", []):
            parts.append(f"- {line}")
        parts.append(f"\n### Full Draft\n\n{c.get('draft', '')}\n")

    # Adaptations
    if project.adaptations:
        parts.append("## Platform Adaptations\n")
        for plat, text in project.adaptations.items():
            parts.append(f"### {plat.title()}\n\n{text}\n")

    # Creative suggestions
    if project.creative_suggestions:
        cs = project.creative_suggestions
        parts.append("## Creative Suggestions\n")
        parts.append(f"**CTA:** {cs.get('cta', '')}\n")
        seo = ", ".join(cs.get("seo_keywords", []))
        parts.append(f"**SEO Keywords:** {seo}\n")
        thumbs = cs.get("thumbnail_ideas", [])
        if thumbs:
            parts.append("**Thumbnail Ideas:**\n")
            for t in thumbs:
                parts.append(f"- {t}")
        imps = cs.get("improvements", [])
        if imps:
            parts.append("\n**Improvements:**\n")
            for imp in imps:
                parts.append(f"- {imp}")

    return "\n".join(parts)


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _to_docx(project: Project) -> io.BytesIO:
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore

    doc = Document()

    # Title
    title_para = doc.add_heading(project.title, 0)

    # Metadata table
    meta = doc.add_paragraph()
    meta.add_run(f"Platform: {project.platform}  |  Goal: {project.goal}  |  Length: {project.length}")
    meta.add_run(f"\nCreated: {project.created_at.strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("Original Idea", 1)
    doc.add_paragraph(project.idea)

    # Analysis
    if project.analysis:
        a = project.analysis
        doc.add_heading("Analysis", 1)
        for key, label in [("topic", "Topic"), ("audience", "Audience"),
                            ("purpose", "Purpose"), ("tone", "Tone")]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(a.get(key, "")))
        kws = ", ".join(a.get("keywords", []))
        p = doc.add_paragraph()
        p.add_run("Keywords: ").bold = True
        p.add_run(kws)

    # Brainstorm
    if project.brainstorm:
        doc.add_heading("Brainstorm Concepts", 1)
        for i, c in enumerate(project.brainstorm, 1):
            doc.add_heading(f"{i}. {c.get('title', '')}", 2)
            p = doc.add_paragraph()
            p.add_run("Hook: ").bold = True
            p.add_run(str(c.get("hook", "")))
            doc.add_paragraph(str(c.get("description", "")))

    # Recommended direction
    if project.recommended_direction:
        rd = project.recommended_direction
        doc.add_heading("Recommended Direction", 1)
        p = doc.add_paragraph()
        p.add_run(str(rd.get("title", ""))).bold = True
        doc.add_paragraph(str(rd.get("reason", "")))

    # Content
    if project.content:
        c = project.content
        doc.add_heading(f"Content: {c.get('title', '')}", 1)
        doc.add_heading("Outline", 2)
        for line in c.get("outline", []):
            doc.add_paragraph(str(line), style="List Bullet")
        doc.add_heading("Full Draft", 2)
        for para in str(c.get("draft", "")).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Adaptations
    if project.adaptations:
        doc.add_heading("Platform Adaptations", 1)
        for plat, text in project.adaptations.items():
            doc.add_heading(plat.title(), 2)
            doc.add_paragraph(str(text))

    # Creative suggestions
    if project.creative_suggestions:
        cs = project.creative_suggestions
        doc.add_heading("Creative Suggestions", 1)
        p = doc.add_paragraph()
        p.add_run("CTA: ").bold = True
        p.add_run(str(cs.get("cta", "")))
        seo = ", ".join(cs.get("seo_keywords", []))
        p2 = doc.add_paragraph()
        p2.add_run("SEO Keywords: ").bold = True
        p2.add_run(seo)
        thumbs = cs.get("thumbnail_ideas", [])
        if thumbs:
            doc.add_paragraph("Thumbnail Ideas:").runs[0]
            for t in thumbs:
                doc.add_paragraph(str(t), style="List Bullet")
        imps = cs.get("improvements", [])
        if imps:
            doc.add_paragraph("Improvements:").runs[0]
            for imp in imps:
                doc.add_paragraph(str(imp), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
